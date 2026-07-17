"""
docx_parser.py
Lectura de plantillas Word (.docx), detección automática de campos
{campo} y generación de cartas individuales reemplazando esos campos.
"""

import os
import re
from docx import Document

PLACEHOLDER_PATTERN = re.compile(r"\{([^{}]+)\}")


def _extract_text_from_doc(doc):
    """Concatena todo el texto visible de un documento (párrafos y tablas)."""
    texts = []
    for p in doc.paragraphs:
        texts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    texts.append(p.text)
    return "\n".join(texts)


def detect_fields(template_path):
    """Devuelve la lista ordenada (sin duplicados) de nombres de campo
    encontrados en la plantilla, ej: ['empresa', 'monto', 'numero_cuenta'].
    Busca {nombre_campo}, «nombre_campo» y <<nombre_campo>> tanto en párrafos como en tablas."""
    doc = Document(template_path)
    full_text = _extract_text_from_doc(doc)
    
    # Detectar placeholders tipo {campo}, chevrons «campo» y <<campo>>
    found_braces = re.findall(r"\{([^{}]+)\}", full_text)
    found_chevrons = re.findall(r"«([^«»]+)»", full_text)
    found_double_less = re.findall(r"<<([^<>]+)>>", full_text)
    
    seen = []
    for f in found_braces + found_chevrons + found_double_less:
        name = f.strip()
        if name and name not in seen:
            seen.append(name)
    return seen


def _replace_in_paragraph(paragraph, values):
    """Reemplaza placeholders {campo}, «campo» y <<campo>> en un párrafo conservando el formato
    lo mejor posible."""
    full_text = "".join(run.text for run in paragraph.runs)
    if "{" not in full_text and "«" not in full_text and "<" not in full_text:
        return
    new_text = full_text
    for key, val in values.items():
        val_str = "" if val is None else str(val)
        new_text = new_text.replace("{" + key + "}", val_str)
        new_text = new_text.replace("«" + key + "»", val_str)
        new_text = new_text.replace("<<" + key + ">>", val_str)
    if new_text == full_text:
        return
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def generate_letter(template_path, values, output_path):
    """Genera un .docx a partir de la plantilla, reemplazando los
    placeholders con los valores del diccionario `values`.
    Devuelve la ruta del archivo generado."""
    doc = Document(template_path)

    for p in doc.paragraphs:
        _replace_in_paragraph(p, values)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, values)

    out_dir = os.path.dirname(output_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    doc.save(output_path)
    return output_path
